from pathlib import Path
from copy import deepcopy
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
RAW = ROOT / "docs/user-manual/screenshots/raw"
IMAGES = ROOT / "docs/user-manual/screenshots/final"
V2_RAW = ROOT / "docs/user-manual/screenshots/v2-raw"
V2_IMAGES = ROOT / "docs/user-manual/screenshots/v2-final"
OUTPUT = ROOT / "output/user-manuals"
IMAGES.mkdir(parents=True, exist_ok=True)
V2_IMAGES.mkdir(parents=True, exist_ok=True)
OUTPUT.mkdir(parents=True, exist_ok=True)

PINK = "EE2B5B"
DARK = "252A34"
MUTED = "667085"
LIGHT = "F7F8FA"
PALE_PINK = "FFF1F4"
WHITE = "FFFFFF"


def sanitize_screenshots():
    for src in RAW.glob("*.png"):
        img = Image.open(src).convert("RGB")
        # Remove Brave tabs, address bar, bookmarks, and all unrelated browser UI.
        app = img.crop((0, 84, img.width, img.height))
        # Remove the documentation account name/role from the app header.
        draw = ImageDraw.Draw(app)
        draw.rectangle((1230, 25, app.width, 63), fill=(255, 255, 255))
        # Some React tables briefly retained their previous rows while the
        # privacy-safe search was being applied. Keep the real controls and
        # column headings, but remove all preview-record rows from list views.
        if src.stem in {
            "en-simulations", "en-clients", "en-users", "en-agencies",
            "en-base-values", "es-simulations",
        }:
            draw.rectangle((187, 171, app.width - 14, 632), fill=(255, 255, 255))
        app.save(IMAGES / f"{src.stem}.jpg", quality=92, optimize=True)

    # Detailed simulation captures. These retain real layout and controls while
    # opaque privacy panels remove customer, supply-point and invoice values.
    for src in V2_RAW.glob("*.png"):
        img = Image.open(src).convert("RGB")
        # Remove browser chrome and the Sys Admin sidebar. The detailed
        # simulation workflow is shared by roles, while sidebar access is not.
        app = img.crop((180, 84, img.width, img.height))
        draw = ImageDraw.Draw(app)
        draw.rectangle((1050, 0, app.width, 48), fill=(255, 255, 255))
        stem = src.stem.removeprefix("es-")
        if stem in {"sim-extracted", "sim-validated"}:
            # Preserve headings, progress indicators and validation controls;
            # hide all invoice values in the form body.
            draw.rectangle((25, 248, app.width - 28, app.height - 22), fill=(247, 248, 250))
            draw.text((45, 268), "Example invoice values hidden for privacy", fill=(102, 112, 133))
        elif stem == "sim-inputs":
            draw.rectangle((8, 53, app.width - 18, 84), fill=(255, 255, 255))
            draw.rectangle((25, 243, app.width - 40, 444), fill=(247, 248, 250))
            draw.text((45, 266), "Customer and supply values hidden for privacy", fill=(102, 112, 133))
        elif stem == "sim-results":
            draw.rectangle((8, 53, app.width - 18, 84), fill=(255, 255, 255))
        app.save(V2_IMAGES / f"{stem}.jpg", quality=92, optimize=True)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, bold=False, color=DARK, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run(run, 9, color=MUTED)


def configure_document(doc, role, language):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.42)
    sec.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, before, after in (
        ("Heading 1", 17, 18, 8),
        ("Heading 2", 13.5, 14, 6),
        ("Heading 3", 11.5, 10, 4),
    ):
        st = styles[name]
        st.font.name = "Aptos Display"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(PINK)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        st = styles[list_name]
        st.font.name = "Aptos"
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.18

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("AXPO SIMULATOR  |  " + role.upper())
    set_run(r, 8.5, bold=True, color=PINK)

    footer = sec.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.columns[0].width = Inches(5.6)
    table.columns[1].width = Inches(0.9)
    left = table.cell(0, 0).paragraphs[0]
    left.text = "User Manual | Preview UI reviewed 17 August 2026" if language == "en" else "Manual de usuario | Interfaz de preproducción revisada el 17 de agosto de 2026"
    set_run(left.runs[0], 8, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(right)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)


def add_cover(doc, role, language):
    is_en = language == "en"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AXPO")
    set_run(r, 18, bold=True, color=PINK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("AXPO Simulator")
    set_run(r, 30, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(("User Manual" if is_en else "Manual de usuario") + f" - {role}")
    set_run(r, 18, bold=True, color=PINK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("Create, review and share energy offers from one workspace." if is_en else "Crea, revisa y comparte ofertas de energía desde un único espacio de trabajo.")
    set_run(r, 12, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("Version 1.0 | 17 August 2026" if is_en else "Versión 1.0 | 17 de agosto de 2026")
    set_run(r, 9.5, color=MUTED)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r)


def add_callout(doc, label, text, kind="tip"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.42)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_PINK if kind != "warning" else "FFF6E5")
    set_cell_margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    r = p.add_run(label + " ")
    set_run(r, 10, bold=True, color=PINK if kind != "warning" else "9A6700")
    r = p.add_run(text)
    set_run(r, 10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc, language, name, caption):
    path = V2_IMAGES / f"{name}.jpg" if name.startswith("sim-") else IMAGES / f"{language}-{name}.jpg"
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.45))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(8)
    cp.paragraph_format.keep_with_next = False
    r = cp.add_run(caption)
    set_run(r, 8.5, italic=True, color=MUTED)


def add_capabilities_table(doc, role, language):
    en = language == "en"
    rows = {
        "Administrator": [
            ("Simulations", "Create, review, edit, share, duplicate, archive, download and delete."),
            ("Clients", "Create, edit, deactivate, restore and delete."),
            ("Users & agencies", "Create and maintain accounts, roles, sessions, access requests and agencies."),
            ("Base values", "Import, review, activate and archive economic reference sets."),
            ("Operations", "Analytics, audit logs, simulation issues, OCR usage and system configuration."),
        ],
        "Agent": [
            ("Simulations", "Create, review, edit, share, duplicate and archive."),
            ("Clients", "View, create, edit and delete client records."),
            ("Analytics", "View operational metrics within the permitted scope."),
            ("Restricted", "No base values, users, agencies, logs, OCR usage or configuration administration."),
        ],
        "Commercial": [
            ("Simulations", "Create, review, edit, share, duplicate and archive."),
            ("Clients", "View, create, edit and delete client records."),
            ("Restricted", "No base values, analytics, users, agencies, logs, OCR usage or configuration."),
        ],
    }
    es_rows = {
        "Administrador": [
            ("Simulaciones", "Crear, revisar, editar, compartir, duplicar, archivar, descargar y eliminar."),
            ("Clientes", "Crear, editar, desactivar, restaurar y eliminar."),
            ("Usuarios y agencias", "Crear y mantener cuentas, roles, sesiones, solicitudes de alta y agencias."),
            ("Valores base", "Importar, revisar, activar y archivar conjuntos de referencia económica."),
            ("Operaciones", "Analíticas, auditoría, incidencias de simulación, uso OCR y configuración."),
        ],
        "Agente": [
            ("Simulaciones", "Crear, revisar, editar, compartir, duplicar y archivar."),
            ("Clientes", "Consultar, crear, editar y eliminar registros de clientes."),
            ("Analíticas", "Consultar métricas operativas dentro del ámbito permitido."),
            ("Restringido", "Sin valores base, usuarios, agencias, registros, uso OCR o configuración."),
        ],
        "Comercial": [
            ("Simulaciones", "Crear, revisar, editar, compartir, duplicar y archivar."),
            ("Clientes", "Consultar, crear, editar y eliminar registros de clientes."),
            ("Restringido", "Sin valores base, analíticas, usuarios, agencias, registros, uso OCR o configuración."),
        ],
    }
    es_key = {"Administrator": "Administrador", "Agent": "Agente", "Commercial": "Comercial"}[role]
    data = rows[role] if en else es_rows[es_key]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.6)
    table.columns[1].width = Inches(4.8)
    hdr = table.rows[0].cells
    hdr[0].text = "Area" if en else "Área"
    hdr[1].text = "Access for this role" if en else "Acceso para este rol"
    hdr[0].width = Inches(1.6)
    hdr[1].width = Inches(4.8)
    set_repeat_table_header(table.rows[0])
    for c in hdr:
        set_cell_shading(c, PINK)
        set_cell_margins(c)
        for r in c.paragraphs[0].runs:
            set_run(r, 9.5, bold=True, color=WHITE)
    for label, desc in data:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = desc
        cells[0].width = Inches(1.6)
        cells[1].width = Inches(4.8)
        for i, c in enumerate(cells):
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c)
            if i == 0:
                set_cell_shading(c, LIGHT)
            for r in c.paragraphs[0].runs:
                set_run(r, 9.3, bold=(i == 0))


COMMON_EN = {
    "contents": ["Purpose and role scope", "Signing in securely", "Application navigation", "Working with tables", "Creating a simulation", "Reviewing and sharing offers", "Help, preferences and troubleshooting"],
    "intro": "This manual explains the tasks available to the selected role in AXPO Simulator. Screens were captured from the preview environment. Available records and labels can change as the application evolves or when permissions are customized.",
    "login_steps": [
        "Open the AXPO Simulator internal address provided by your organization.",
        "Enter your work email address and password, then select Sign in.",
        "When Two-Factor Authentication appears, enter the six-digit code sent to your email.",
        "Select Verify Code. The application opens the Simulations workspace.",
    ],
    "navigation": [
        "The left sidebar contains only the sections available to your role.",
        "The top bar contains the page title, current actions, notifications and the profile menu.",
        "Use global search (Command+K on macOS) to find accessible records and destinations quickly.",
        "Use Help & Tutorials to launch guided walkthroughs that are automatically filtered by permission.",
    ],
    "tables": [
        "Search narrows visible records using free text.",
        "Filters narrow structured fields such as status, agency, date or commodity.",
        "View presets restore useful filter and sorting combinations.",
        "Show/hide columns and Density change the table layout for the current browser.",
        "Select a column header to sort. Select row checkboxes only when a permitted bulk action is required.",
    ],
    "create_steps": [
        "From Simulations, select New simulation.",
        "Upload the complete invoice as PDF or supported image files. For image invoices, include every page.",
        "Confirm the detected invoice provider and invoice type. Correct the provider when detection is wrong.",
        "Select Extract data and wait for OCR processing to finish.",
        "Select an existing client. If your role permits client creation and no match exists, create the client from the dialog.",
        "Compare every extracted value with the original invoice: commodity, CUPS, dates, consumption, contracted power, prices, taxes and total.",
        "Correct inaccurate or missing values, then select Validate.",
        "Complete the remaining offer fields and expiration date. Select Create simulation only after the review is complete.",
    ],
    "review_steps": [
        "Find the simulation by reference, client, CUPS, owner, agency, commodity or status.",
        "Select Simulate for a draft or View for a shared simulation.",
        "Check the summary, attached invoice, status, PIN and expiration date.",
        "Use Results to compare products, totals, savings, percentages, charts and historical values.",
        "Use Inputs to review the full set of invoice, client, consumption, power, pricing and tax values.",
        "After changing an input, recalculate and review every result before selecting an offer.",
    ],
    "share_steps": [
        "Select the intended offer after checking totals, assumptions and savings.",
        "Open Share and choose PDF download or email, according to your permission.",
        "Choose the correct language and template. Review client details, selected product, values and branding.",
        "Download the PDF or send the email only when the preview is correct. Sharing changes the simulation status.",
    ],
    "troubleshooting": [
        "No OTP email: wait a minute, check spam, then use Resend when available. Only the newest code is valid.",
        "Create button disabled: review required fields and complete validation.",
        "Unexpected OCR value: correct it before validation and use Report issue when the extraction should be investigated.",
        "Missing action or section: the capability may not be granted to your role. Contact an administrator.",
        "Outdated screen after an update: refresh the page. If the issue continues, sign out and sign in again.",
    ],
}

COMMON_ES = {
    "contents": ["Objetivo y alcance del rol", "Inicio de sesión seguro", "Navegación de la aplicación", "Trabajo con tablas", "Creación de una simulación", "Revisión y envío de ofertas", "Ayuda, preferencias y solución de problemas"],
    "intro": "Este manual explica las tareas disponibles para el rol seleccionado en AXPO Simulator. Las capturas proceden del entorno de preproducción. Los registros, etiquetas y permisos pueden cambiar a medida que evoluciona la aplicación.",
    "login_steps": [
        "Abre la dirección interna de AXPO Simulator facilitada por tu organización.",
        "Introduce tu correo corporativo y contraseña y selecciona Iniciar sesión.",
        "Cuando aparezca la autenticación en dos pasos, introduce el código de seis dígitos enviado por correo.",
        "Selecciona Verificar código. La aplicación abrirá el área Simulaciones.",
    ],
    "navigation": [
        "La barra lateral muestra únicamente las secciones disponibles para tu rol.",
        "La barra superior contiene el título, las acciones de la página, notificaciones y el menú de perfil.",
        "Utiliza la búsqueda global (Comando+K en macOS) para localizar destinos y registros accesibles.",
        "En Ayuda y tutoriales encontrarás recorridos guiados filtrados automáticamente por permisos.",
    ],
    "tables": [
        "Buscar reduce los registros visibles mediante texto libre.",
        "Filtros limita campos estructurados como estado, agencia, fecha o tipo de energía.",
        "Las vistas guardadas recuperan combinaciones útiles de filtros y ordenación.",
        "Mostrar/ocultar columnas y Densidad adaptan la tabla en el navegador actual.",
        "Selecciona una cabecera para ordenar. Marca filas solo cuando necesites una acción masiva permitida.",
    ],
    "create_steps": [
        "En Simulaciones, selecciona Nueva simulación.",
        "Sube la factura completa en PDF o imágenes compatibles. Incluye todas las páginas si son imágenes.",
        "Confirma el proveedor y el tipo de factura detectados. Corrige el proveedor si la detección es incorrecta.",
        "Selecciona Extraer datos y espera a que termine el procesamiento OCR.",
        "Selecciona un cliente existente. Si tu rol permite crear clientes y no hay coincidencias, créalo desde el diálogo.",
        "Compara con la factura original: energía, CUPS, fechas, consumo, potencia contratada, precios, impuestos y total.",
        "Corrige los valores incorrectos o ausentes y selecciona Validar.",
        "Completa los campos restantes y la fecha de caducidad. Selecciona Crear simulación únicamente tras revisar todo.",
    ],
    "review_steps": [
        "Busca la simulación por referencia, cliente, CUPS, propietario, agencia, energía o estado.",
        "Selecciona Simular para un borrador o Ver para una simulación compartida.",
        "Comprueba el resumen, la factura adjunta, el estado, PIN y fecha de caducidad.",
        "En Resultados compara productos, importes, ahorro, porcentajes, gráficos e históricos.",
        "En Datos de entrada revisa todos los valores de factura, cliente, consumo, potencia, precios e impuestos.",
        "Después de cambiar un dato, recalcula y revisa todos los resultados antes de elegir una oferta.",
    ],
    "share_steps": [
        "Selecciona la oferta adecuada tras comprobar importes, hipótesis y ahorro.",
        "Abre Compartir y elige descargar PDF o enviar por correo, según tus permisos.",
        "Selecciona idioma y plantilla. Revisa cliente, producto, valores y marca.",
        "Descarga o envía únicamente cuando la vista previa sea correcta. Compartir cambia el estado de la simulación.",
    ],
    "troubleshooting": [
        "No llega el OTP: espera un minuto, revisa spam y utiliza Reenviar cuando esté disponible. Solo vale el último código.",
        "El botón Crear está desactivado: revisa los campos obligatorios y completa la validación.",
        "Valor OCR incorrecto: corrígelo antes de validar e informa de la incidencia si debe investigarse.",
        "Falta una acción o sección: puede no estar concedida a tu rol. Contacta con un administrador.",
        "Pantalla desactualizada: actualiza la página. Si continúa, cierra la sesión y vuelve a entrar.",
    ],
}


ADMIN_EN = [
    ("Managing clients", "Clients are shared business records used by simulations. Create a client only after checking that it does not already exist.", ["Use search and filters to check for duplicates.", "Select New client and enter company, tax, contact, address, agency and preferred language data.", "Save only verified information. Deactivate records that should no longer be used; restore them when appropriate.", "Permanent deletion is irreversible and should follow the organization's retention policy."], "clients", "Clients list with a privacy-safe empty result", "new-client", "Blank client creation form"),
    ("Managing users", "User administration controls identity, role, agency, preferences and active access.", ["Select New user and enter the identity and contact fields.", "Assign the minimum role and correct agency required for the person's work.", "Review Preferences before saving, then confirm the welcome/setup process.", "Use Access requests to verify and approve or reject requested access.", "Use Sessions to inspect or revoke active sessions; deactivation prevents future sign-in."], "users", "User list, search and access actions", "new-user", "Blank user form and role assignment"),
    ("Managing agencies", "Agencies scope users, clients, commercial rules and product availability.", ["Create an agency with a clear unique name and accurate location.", "Enable TLV only for the appropriate agency type.", "In the agency detail, review users, products, tariffs and operational settings.", "Test a sample simulation after tariff or product changes."], "agencies", "Agency list and management actions", "new-agency", "Blank agency creation form"),
    ("Base values", "Base values are economic reference data used by simulation calculations. Versions preserve calculation history.", ["Upload a prepared XLSX or XLSM workbook.", "Choose global or agency scope and verify the target agency.", "Review validation messages, item coverage, version and effective dates.", "Activate only after validation and test a representative simulation.", "Archive superseded versions instead of deleting historical data."], "base-values", "Base-value versions and publishing actions", None, None),
    ("Analytics", "Analytics summarizes simulation creation, sharing, opening and agency performance for the selected period and scope.", ["Select commodity, agency and reporting period.", "Compare like-for-like periods and note the record scope.", "Use counts and date ranges when communicating a result.", "Treat incomplete, archived or expired simulations as possible context for a change."], "analytics", "Analytics dashboard and engagement funnel", None, None),
    ("Audit logs and simulation issues", "The current Administrator configuration provides access to audit activity and reported simulation issues. Email, cron, OCR and application-error logs are not enabled for this role.", ["Filter audit activity by date, type, user or reference before opening details.", "Use the simulation reference and event time to connect a reported issue with relevant audit actions.", "Do not copy sensitive payloads into general communication channels.", "Resolve or dismiss simulation issues only after recording the reason and outcome."], None, None, None, None),
    ("Configurations", "Configuration changes can affect every user, calculation, document or outbound message.", ["Read the current value and description before editing.", "Change one logical setting at a time and keep a rollback value.", "Test changes to PDF/email templates, calculations, OCR prompts and role permissions in preview.", "Role-permission changes should be validated with the affected role before release."], "configurations", "Configuration groups for documents, business rules, platform, AI/OCR and access", None, None),
]

ADMIN_ES = [
    ("Gestión de clientes", "Los clientes son registros compartidos utilizados por las simulaciones. Comprueba primero que no exista el cliente.", ["Busca y filtra para evitar duplicados.", "Selecciona Nuevo cliente e introduce empresa, datos fiscales, contacto, dirección, agencia e idioma.", "Guarda únicamente información verificada. Desactiva registros que ya no deban utilizarse y restáuralos cuando proceda.", "La eliminación permanente es irreversible y debe respetar la política de conservación."], "clients", "Lista de clientes sin datos sensibles", "new-client", "Formulario de alta de cliente vacío"),
    ("Gestión de usuarios", "La administración de usuarios controla identidad, rol, agencia, preferencias y acceso activo.", ["Selecciona Nuevo usuario y completa los datos de identidad y contacto.", "Asigna el rol mínimo y la agencia correcta para el trabajo de la persona.", "Revisa Preferencias antes de guardar y confirma el proceso de bienvenida/alta.", "En Solicitudes de alta, verifica y aprueba o rechaza cada solicitud.", "En Sesiones puedes revisar o revocar accesos activos; la desactivación impide futuros inicios de sesión."], "users", "Lista de usuarios y acciones de acceso", "new-user", "Formulario vacío y asignación de rol"),
    ("Gestión de agencias", "Las agencias delimitan usuarios, clientes, reglas comerciales y disponibilidad de productos.", ["Crea la agencia con un nombre único y ubicación correcta.", "Activa TLV solo para el tipo de agencia adecuado.", "En el detalle revisa usuarios, productos, tarifas y ajustes operativos.", "Prueba una simulación de muestra después de modificar tarifas o productos."], "agencies", "Lista y acciones de agencias", "new-agency", "Formulario de alta de agencia vacío"),
    ("Valores base", "Los valores base son referencias económicas utilizadas por los cálculos. Las versiones conservan el histórico.", ["Sube un libro XLSX o XLSM preparado.", "Elige ámbito global o de agencia y verifica la agencia de destino.", "Revisa mensajes de validación, cobertura, versión y vigencia.", "Activa solo después de validar y probar una simulación representativa.", "Archiva versiones sustituidas en lugar de borrar el histórico."], "base-values", "Versiones y acciones de publicación de valores base", None, None),
    ("Analíticas", "Las analíticas resumen creación, envío, apertura y rendimiento de agencias para el periodo y ámbito seleccionados.", ["Selecciona energía, agencia y periodo.", "Compara periodos equivalentes y anota el ámbito de los registros.", "Incluye recuentos y fechas al comunicar resultados.", "Ten en cuenta simulaciones incompletas, archivadas o caducadas al interpretar cambios."], "analytics", "Panel de analíticas y embudo de actividad", None, None),
    ("Auditoría e incidencias de simulación", "La configuración actual del rol Administrador permite consultar auditoría e incidencias de simulación. Los registros de correo, tareas programadas, OCR y errores de aplicación no están habilitados.", ["Filtra la auditoría por fecha, tipo, usuario o referencia antes de abrir detalles.", "Utiliza la referencia y la hora para relacionar una incidencia con las acciones auditadas.", "No copies cargas sensibles en canales generales de comunicación.", "Resuelve o descarta una incidencia solo después de documentar motivo y resultado."], None, None, None, None),
    ("Configuraciones", "Los cambios de configuración pueden afectar a todos los usuarios, cálculos, documentos o mensajes.", ["Lee el valor actual y su descripción antes de editar.", "Cambia un bloque lógico cada vez y conserva el valor anterior.", "Prueba en preproducción plantillas, cálculos, prompts OCR y permisos.", "Valida los cambios de permisos con el rol afectado antes de publicarlos."], "configurations", "Grupos de configuración de documentos, negocio, plataforma, IA/OCR y acceso", None, None),
]

AGENT_EN = [
    ("Client management", "Agents can view, create, edit and delete client records required by simulation work.", ["Search before creating a client to avoid duplicates.", "Complete company, contact, address, agency and language information.", "Keep tax and contact data aligned with the source invoice.", "Before deleting, confirm the record is the intended client and follow the organization's retention rules."], None, None, None, None),
    ("Analytics", "The dashboard shows the activity and performance available to your operational scope.", ["Select a consistent period and commodity.", "Use totals and date ranges when sharing results.", "Ask an administrator when agency-level visibility appears incomplete."], None, None, None, None),
]

AGENT_ES = [
    ("Gestión de clientes", "Los agentes pueden consultar, crear, editar y eliminar clientes necesarios para las simulaciones.", ["Busca antes de crear para evitar duplicados.", "Completa empresa, contacto, dirección, agencia e idioma.", "Mantén los datos fiscales y de contacto coherentes con la factura.", "Antes de eliminar, confirma el registro y aplica las normas de conservación de la organización."], None, None, None, None),
    ("Analíticas", "El panel muestra la actividad y rendimiento disponibles para tu ámbito operativo.", ["Selecciona un periodo y energía coherentes.", "Incluye totales y rango de fechas al compartir resultados.", "Consulta al administrador si la visibilidad por agencia parece incompleta."], None, None, None, None),
]


def build_manual(role, language):
    en = language == "en"
    common = COMMON_EN if en else COMMON_ES
    role_es = {"Administrator": "Administrador", "Agent": "Agente", "Commercial": "Comercial"}[role]
    display_role = role if en else role_es
    doc = Document()
    configure_document(doc, display_role, language)
    add_cover(doc, display_role, language)

    add_heading(doc, "About this manual" if en else "Acerca de este manual", 1)
    add_body(doc, common["intro"])
    add_callout(doc, "Role scope:" if en else "Ámbito del rol:", "This manual deliberately excludes System Administrator functions." if en else "Este manual excluye deliberadamente las funciones de Administrador del sistema.")

    add_heading(doc, "Contents" if en else "Contenido", 2)
    for item in common["contents"]:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(item))
    if role == "Administrator":
        extra = ["Client, user and agency administration", "Base values, analytics, logs and configurations"] if en else ["Administración de clientes, usuarios y agencias", "Valores base, analíticas, registros y configuraciones"]
        add_bullets(doc, extra)
    elif role == "Agent":
        add_bullets(doc, ["Client management and analytics"] if en else ["Gestión de clientes y analíticas"])

    add_heading(doc, "Role capabilities" if en else "Capacidades del rol", 1)
    add_capabilities_table(doc, role, language)
    add_callout(doc, "Important:" if en else "Importante:", "Permissions can be customized by an administrator. The live application is authoritative when a button or section differs from this guide." if en else "Un administrador puede personalizar los permisos. La aplicación en uso es la referencia si un botón o sección difiere de esta guía.", "warning")

    add_heading(doc, "1. Sign in securely" if en else "1. Iniciar sesión de forma segura", 1)
    add_steps(doc, common["login_steps"])
    add_callout(doc, "Security:" if en else "Seguridad:", "Never share your password, OTP, simulation PIN or client link. Sign out on shared devices." if en else "No compartas contraseña, OTP, PIN de simulación ni enlace de cliente. Cierra la sesión en dispositivos compartidos.", "warning")

    add_heading(doc, "2. Navigate the application" if en else "2. Navegar por la aplicación", 1)
    add_bullets(doc, common["navigation"])
    add_figure(doc, language, "tutorials", "Help & Tutorials presents permission-aware walkthroughs." if en else "Ayuda y tutoriales ofrece recorridos adaptados a los permisos.")

    add_heading(doc, "3. Work with lists and tables" if en else "3. Trabajar con listas y tablas", 1)
    add_bullets(doc, common["tables"])
    add_figure(doc, language, "simulations", "The simulations list with search, filters, columns, density and row actions." if en else "Lista de simulaciones con búsqueda, filtros, columnas, densidad y acciones.")
    add_callout(doc, "Privacy:" if en else "Privacidad:", "Clear searches and filters before sharing your screen. Do not place client or user data in screenshots or support messages unless required and authorized." if en else "Limpia búsquedas y filtros antes de compartir pantalla. No incluyas datos de clientes o usuarios en capturas o mensajes salvo necesidad y autorización.")

    add_heading(doc, "4. Create a simulation: invoice to validated data" if en else "4. Crear una simulación: de la factura a los datos validados", 1)
    add_body(doc, "Simulation work follows four visible stages: Source, Extract, Validate and Create. Finish each stage in order; the progress strip shows what is ready and what still needs attention." if en else "El trabajo de simulación sigue cuatro etapas visibles: Origen, Extraer, Validar y Crear. Completa cada etapa en orden; la barra de progreso indica qué está listo y qué requiere atención.")
    add_heading(doc, "4.1 Upload the source invoice" if en else "4.1 Subir la factura de origen", 2)
    add_figure(doc, language, "sim-source", "Start with the complete invoice and the correct customer, commodity and expiry date." if en else "Empieza con la factura completa y confirma cliente, energía y fecha de caducidad.")
    add_steps(doc, common["create_steps"][:3])
    add_figure(doc, language, "sim-provider", "Provider recognition must be confirmed before OCR extraction begins." if en else "Confirma el proveedor reconocido antes de iniciar la extracción OCR.")
    add_heading(doc, "4.2 Extract the invoice" if en else "4.2 Extraer la factura", 2)
    add_steps(doc, common["create_steps"][3:5])
    add_figure(doc, language, "sim-extracting", "Keep the page open while OCR reads and structures the invoice." if en else "Mantén la página abierta mientras el OCR lee y estructura la factura.")
    add_heading(doc, "4.3 Review and validate every field" if en else "4.3 Revisar y validar todos los campos", 2)
    add_steps(doc, common["create_steps"][5:])
    add_figure(doc, language, "sim-validated", "Validation confirms that the reviewed form can proceed to simulation creation; example values are hidden." if en else "La validación confirma que el formulario revisado puede pasar a la creación; los valores de ejemplo están ocultos.")
    add_callout(doc, "Quality check:" if en else "Control de calidad:", "OCR accelerates entry but does not replace verification against the source invoice." if en else "El OCR acelera la carga, pero no sustituye la comparación con la factura original.", "warning")

    add_heading(doc, "5. Review, adjust and compare a simulation" if en else "5. Revisar, ajustar y comparar una simulación", 1)
    add_body(doc, "The simulation detail is the operational centre of AXPO Simulator. Its header identifies the reference, status, expiry, invoice and PIN; the Data and Results tabs separate calculation inputs from commercial output." if en else "El detalle es el centro operativo de AXPO Simulator. La cabecera identifica referencia, estado, caducidad, factura y PIN; las pestañas Datos y Resultados separan entradas de cálculo y salida comercial.")
    add_steps(doc, common["review_steps"])
    add_heading(doc, "5.1 Inspect calculation inputs" if en else "5.1 Revisar los datos de cálculo", 2)
    add_figure(doc, language, "sim-inputs", "Inputs group customer, supply, billing, consumption, power, charges and tax data; identifying values are hidden." if en else "Datos agrupa cliente, suministro, facturación, consumo, potencia, cargos e impuestos; los identificadores están ocultos.")
    add_callout(doc, "Recalculation:" if en else "Recálculo:", "Any changed input can alter product order, savings and totals. Select Recalculate offers and repeat the comparison before sharing." if en else "Cualquier cambio puede alterar el orden, el ahorro y los totales. Selecciona Recalcular ofertas y repite la comparación antes de compartir.", "warning")
    add_heading(doc, "5.2 Compare offers" if en else "5.2 Comparar ofertas", 2)
    add_figure(doc, language, "sim-results", "Filter fixed, indexed and customized products, then compare AXPO total, monthly saving, percentage difference and annual saving." if en else "Filtra productos fijos, indexados y personalizados y compara total AXPO, ahorro mensual, diferencia porcentual y ahorro anual.")
    add_bullets(doc, [
        "The Best offer marker identifies the lowest calculated total; it is not an automatic recommendation.",
        "A positive saving should still be checked against product conditions, billing month and all assumptions.",
        "Select only one product after confirming that its commercial conditions fit the customer.",
    ] if en else [
        "La marca Mejor oferta identifica el menor total calculado; no es una recomendación automática.",
        "Aunque exista ahorro, comprueba condiciones, mes de facturación y todas las hipótesis.",
        "Selecciona un único producto tras confirmar que sus condiciones encajan con el cliente.",
    ])
    add_heading(doc, "5.3 Select and share an offer" if en else "5.3 Seleccionar y compartir una oferta", 2)
    add_steps(doc, common["share_steps"])
    add_callout(doc, "Before sharing:" if en else "Antes de compartir:", "Confirm the client, selected product, language, recipient, totals, savings and expiration date. Email and public-link actions affect external recipients." if en else "Confirma cliente, producto, idioma, destinatario, importes, ahorro y caducidad. Las acciones de correo y enlace afectan a destinatarios externos.", "warning")

    if role == "Administrator":
        sections = ADMIN_EN if en else ADMIN_ES
    elif role == "Agent":
        sections = AGENT_EN if en else AGENT_ES
    else:
        sections = []

    section_no = 6
    for title, intro, bullets, image1, caption1, image2, caption2 in sections:
        add_heading(doc, f"{section_no}. {title}", 1)
        add_body(doc, intro)
        if image1:
            add_figure(doc, language, image1, caption1)
        add_steps(doc, bullets)
        if image2:
            add_figure(doc, language, image2, caption2)
        section_no += 1

    if role == "Commercial":
        add_heading(doc, "6. Client management" if en else "6. Gestión de clientes", 1)
        add_body(doc, "Commercial users can view, create, edit and delete client records used by simulations." if en else "Los usuarios Comerciales pueden consultar, crear, editar y eliminar los clientes utilizados por las simulaciones.")
        add_bullets(doc, [
            "Search before creating a client to avoid duplicates.",
            "Keep company, tax, contact, address, agency and language information accurate.",
            "Confirm the intended record and retention requirements before deleting a client.",
            "Base values, analytics, users, agencies, logs, OCR usage and configuration remain unavailable.",
        ] if en else [
            "Busca antes de crear para evitar duplicados.",
            "Mantén correctos los datos de empresa, fiscales, contacto, dirección, agencia e idioma.",
            "Confirma el registro y las normas de conservación antes de eliminar un cliente.",
            "Valores base, analíticas, usuarios, agencias, registros, uso OCR y configuración siguen sin estar disponibles.",
        ])
        section_no = 7

    add_heading(doc, f"{section_no}. Profile, language and help" if en else f"{section_no}. Perfil, idioma y ayuda", 1)
    add_bullets(doc, [
        "Open My profile to review your information and personal defaults.",
        "Use the profile menu language selector to switch the interface language.",
        "Use Notifications for application updates that require your attention.",
        "Open Help & Tutorials to repeat guided workflows at any time.",
    ] if en else [
        "Abre Mi perfil para revisar tus datos y valores personales.",
        "Utiliza el selector de idioma del menú de perfil para cambiar la interfaz.",
        "Consulta Notificaciones para avisos que requieran atención.",
        "Abre Ayuda y tutoriales para repetir recorridos guiados cuando lo necesites.",
    ])

    add_heading(doc, f"{section_no + 1}. Troubleshooting" if en else f"{section_no + 1}. Solución de problemas", 1)
    add_bullets(doc, common["troubleshooting"])
    add_callout(doc, "Support request:" if en else "Solicitud de soporte:", "Include the page, simulation reference (when relevant), time, expected result, actual result and a privacy-safe screenshot. Never include passwords, OTP codes or simulation PINs." if en else "Incluye página, referencia (si procede), hora, resultado esperado, resultado real y una captura sin datos sensibles. Nunca incluyas contraseñas, OTP o PIN de simulación.")

    add_heading(doc, "Quick pre-share checklist" if en else "Lista rápida antes de compartir", 1)
    add_bullets(doc, [
        "Correct client and CUPS",
        "Correct commodity, period, consumption and power",
        "Correct pricing, taxes and calculated totals",
        "Intended offer selected",
        "Correct language and template",
        "Correct recipient and expiration date",
        "Preview checked; no internal notes or unintended data",
    ] if en else [
        "Cliente y CUPS correctos",
        "Energía, periodo, consumo y potencia correctos",
        "Precios, impuestos y totales correctos",
        "Oferta prevista seleccionada",
        "Idioma y plantilla correctos",
        "Destinatario y caducidad correctos",
        "Vista previa revisada, sin notas internas ni datos no deseados",
    ])

    slug_role = role.lower()
    filename = f"axpo-simulator-{slug_role}-manual-{language}.docx"
    path = OUTPUT / filename
    doc.save(path)
    return path


def main():
    sanitize_screenshots()
    paths = []
    for language in ("en", "es"):
        for role in ("Administrator", "Agent", "Commercial"):
            paths.append(build_manual(role, language))
    print("\n".join(str(p) for p in paths))


if __name__ == "__main__":
    main()
